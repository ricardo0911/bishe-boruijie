from __future__ import annotations

from pathlib import Path

from docx import Document


ADDED_LINES = [
    "（6）订单明细实体：该实体属性包括明细编号、订单编号、商品编号、商品名称、单价、数量、小计金额、创建时间。具体如图3-10所示。",
    "图3-10 订单明细实体属性描述",
    "（7）库存锁流水实体：该实体属性包括流水编号、订单编号、订单号、花材编号、批次编号、锁定数量、锁定状态、过期时间。具体如图3-11所示。",
    "图3-11 库存锁流水实体属性描述",
    "（8）支付流水实体：该实体属性包括流水编号、订单编号、订单号、支付交易号、支付渠道、支付金额、回调时间、处理结果。具体如图3-12所示。",
    "图3-12 支付流水实体属性描述",
    "（9）推荐结果实体：该实体属性包括结果编号、用户编号、商品编号、推荐分值、推荐理由、生成时间。具体如图3-13所示。",
    "图3-13 推荐结果实体属性描述",
]


def main() -> int:
    docx_path = Path("花店智能管理系统-毕业论文-初稿.docx")
    if not docx_path.exists():
        raise FileNotFoundError(f"未找到文件: {docx_path}")

    doc = Document(docx_path)
    paras = doc.paragraphs

    anchor = None
    for p in paras:
        if (p.text or "").strip().startswith("3.3.2 逻辑结构设计"):
            anchor = p
            break
    if anchor is None:
        raise RuntimeError("未找到锚点段落：3.3.2 逻辑结构设计")

    # Remove old injected block if it exists, then re-insert in exact order.
    remove_prefixes = (
        "（6）订单明细实体",
        "图3-10 订单明细实体属性描述",
        "（7）库存锁流水实体",
        "图3-11 库存锁流水实体属性描述",
        "（8）支付流水实体",
        "图3-12 支付流水实体属性描述",
        "（9）推荐结果实体",
        "图3-13 推荐结果实体属性描述",
    )
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        if any(t.startswith(x) for x in remove_prefixes):
            p._element.getparent().remove(p._element)

    # Re-anchor after deletion.
    anchor = None
    for p in doc.paragraphs:
        if (p.text or "").strip().startswith("3.3.2 逻辑结构设计"):
            anchor = p
            break
    if anchor is None:
        raise RuntimeError("删除后未找到锚点段落：3.3.2 逻辑结构设计")

    for line in ADDED_LINES:
        anchor.insert_paragraph_before(line)

    doc.save(docx_path)
    print("patched", docx_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
