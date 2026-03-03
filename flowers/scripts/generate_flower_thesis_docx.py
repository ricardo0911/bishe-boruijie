from __future__ import annotations

from pathlib import Path

from docx import Document


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def table_with_caption(
    doc: Document, caption: str, headers: list[str], rows: list[list[str]]
) -> None:
    doc.add_paragraph(caption)
    t = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row in rows:
        c = t.add_row().cells
        for i, v in enumerate(row):
            c[i].text = v


def build_doc(output_path: Path) -> None:
    doc = Document()

    p(doc, "青岛农业大学海都学院")
    p(doc, "本科生毕业论文（设计）")
    p(doc, "题目：基于微信小程序的花店智能管理系统设计与实现")
    p(doc, "系别：信息工程系")
    p(doc, "专业：计算机科学与技术")
    p(doc, "班级：计科2201")
    p(doc, "姓名：薄睿洁")
    p(doc, "学号：202200210118")
    p(doc, "指导教师：指导教师")
    p(doc, "2026年2月27日")

    doc.add_page_break()
    p(doc, "基于微信小程序的花店智能管理系统设计与实现")
    p(doc, "计算机科学与技术专业 薄睿洁")
    p(doc, "摘要")
    p(
        doc,
        "鲜花行业具有保鲜周期短、节日需求波动大、库存损耗高等经营特点，传统人工管理方式在订单处理、库存核算和补货决策上容易出现响应滞后与数据不一致问题。针对上述痛点，本文设计并实现了一个基于微信小程序的花店智能管理系统。系统以“用户下单—库存锁定—支付扣减—取消或退款回滚—数据分析与补货建议”为主链路，构建前后端分离与分析服务协同的业务闭环。技术上，系统采用 Spring Boot 提供 REST 接口，MySQL 存储业务数据，微信小程序与 Vue 管理端承载多角色交互，Python 分析服务输出销量预测、补货建议与推荐结果。实现过程中引入 FEFO 批次策略和库存锁机制，保障订单并发场景下的数据一致性。测试结果表明，系统能够稳定完成核心交易流程，库存预警与分析建议可为商家运营提供有效支撑，具有较好的工程可实施性与毕业设计应用价值。"
    )
    p(doc, "关键词：微信小程序；花店管理；FEFO；库存预警；销量预测")

    doc.add_page_break()
    p(doc, "Design and Implementation of a WeChat Mini Program Based Smart Flower Shop Management System")
    p(doc, "Major: Computer Science and Technology  Name: Bo Ruijie")
    p(doc, "Abstract")
    p(
        doc,
        "To address the challenges of high perishability, demand volatility, and inventory inconsistency in flower retail, this thesis designs and implements a smart flower shop management system based on WeChat Mini Program. The system builds a closed-loop workflow including order creation, stock locking, payment confirmation, rollback on cancellation or refund, and data-driven replenishment. The backend is implemented with Spring Boot REST services, the core data is persisted in MySQL, and Python analytics services generate sales forecasting, replenishment suggestions, and recommendation results. A FEFO-based batch strategy and stock-lock mechanism are introduced to improve consistency under concurrent order scenarios. The implementation includes customer-facing mini program pages, merchant operation pages, and admin governance pages. Functional tests show that the system can complete key business processes reliably and provide practical decision support for daily operation. The project demonstrates feasibility, scalability, and engineering value for undergraduate graduation design."
    )
    p(
        doc,
        "Key words: WeChat Mini Program; Flower Shop Management; FEFO; Inventory Alert; Sales Forecasting",
    )

    p(doc, "1 概述")
    p(
        doc,
        "本课题围绕花店经营中的订单效率、库存损耗与运营决策三类核心问题展开。论文在项目真实代码与数据库实现基础上，形成需求分析、系统设计、实现与测试的完整技术闭环，为中小花店数字化提供可落地方案。"
    )

    p(doc, "1.1 课题研究背景和意义")
    p(
        doc,
        "随着移动互联网消费习惯的普及，花店业务已从电话与到店下单逐步迁移到线上渠道。鲜花商品的时效性与易损耗性决定了库存管理必须具备精细化能力，而传统经验式进货与手工记账无法满足高峰期快速响应需求，容易出现超卖、错卖和积压。"
    )
    p(
        doc,
        "本课题通过构建微信小程序用户端、商家管理端和管理员治理端，打通交易链路与库存链路，并将分析结果回流到运营端，能够在工程实践层面验证“业务系统+分析服务”协同模式的可行性，对毕业设计与实际经营均具有应用价值。"
    )

    p(doc, "1.2 国内外研究现状")
    p(
        doc,
        "国外电商管理系统研究较早，重点在于服务化架构、库存优化与数据驱动运营；国内研究更关注行业落地与微信生态应用，围绕小程序交易、库存预警、商家运营报表形成了大量实践案例。"
    )
    p(
        doc,
        "现有研究在通用电商场景中较为成熟，但对鲜花行业“保鲜期约束+批次质量差异+节日波动”的联合处理不足。本文将 FEFO 批次策略、库存锁流水与预测补货结合，形成更贴合花店业务特征的实现路径。"
    )
    p(
        doc,
        "本文在需求分析、架构设计、数据库建模、预测与推荐方法、测试方法等方面综合参考了[1][2][3][4][5][6][7][8][9][10][11][12][13][14][15][16][17][18]。"
    )

    p(doc, "1.3 课题主要研究内容")
    p(
        doc,
        "本课题以“稳定交易链路、可追踪库存、可解释分析建议”为目标，围绕角色分工和业务流程构建系统功能框架。根据项目实现结果，主要研究内容如下。"
    )
    p(doc, "（1）用户端业务模块：实现商品浏览、购物车、下单支付、订单追踪、评价反馈等闭环能力。")
    p(doc, "（2）商家端运营模块：实现商品维护、订单处理、库存预警、销量统计与补货决策查看。")
    p(doc, "（3）管理员治理模块：实现用户与商家管理、系统参数配置、全局订单监管与异常治理。")
    p(doc, "（4）分析决策模块：实现销量预测、推荐计算与补货建议生成，辅助商家运营优化。")

    p(doc, "1.4 论文组织结构")
    p(
        doc,
        "第一章为概述，说明研究背景、意义与研究内容；第二章为需求分析，从可行性、功能性和非功能性三个维度明确系统目标；第三章为概要设计，完成架构、模块与数据库设计；第四章为详细设计与实现，给出关键流程与落地实现；第五章为系统测试，验证核心业务链路；第六章为总结，归纳成果与后续优化方向。"
    )

    p(doc, "2 需求分析")
    p(
        doc,
        "本章结合花店业务场景和现有工程实现，从可行性、功能需求、性能需求与开发环境四个方面展开，为后续设计与实现提供约束依据。"
    )

    p(doc, "2.1 可行性分析")
    p(doc, "系统可行性从经济可行性与技术可行性两个方面进行论证。")

    p(doc, "2.1.1 经济可行性分析")
    p(
        doc,
        "系统采用微信小程序与开源技术栈，部署成本和运维成本可控。后端基于 Spring Boot，前端基于 Vue 与小程序原生框架，数据库采用 MySQL，能够充分利用现有学习与实验环境，降低额外采购投入。"
    )
    p(
        doc,
        "从收益角度看，系统可减少人工对账与库存统计时间，降低缺货和积压损耗，提升订单处理效率和用户复购率，具有明确的经济可行性。"
    )

    p(doc, "2.1.2 技术可行性分析")
    p(
        doc,
        "项目采用分层架构与 REST 接口模式，技术方案成熟、社区支持完善。核心事务在后端通过订单与库存锁联动保证一致性，分析层通过 Python 任务输出可解释结果，能够满足毕业设计“可实现、可演示、可验证”的要求。"
    )
    p(
        doc,
        "从当前仓库代码看，系统已具备用户、商品、订单、库存、售后、统计与分析接口，前端路由与页面结构完整，技术可行性充分。"
    )

    p(doc, "2.2 功能性需求分析")
    p(
        doc,
        "功能性需求围绕用户端、商家端、管理员端三类角色展开，并以订单主链路与库存一致性作为核心约束。"
    )

    p(doc, "2.2.1 系统功能概述")
    p(doc, "系统包括用户端、商家端、管理员端三个角色。")
    p(doc, "用户端功能：商品浏览、分类检索、购物车、下单支付、订单查询、评价反馈。")
    p(doc, "商家端功能：商品管理、订单处理、库存预警、销售报表、补货建议查看。")
    p(doc, "管理员端功能：用户管理、商家管理、订单监管、系统配置与治理。")

    p(doc, "2.2.2 系统功能需求分析")
    p(doc, "系统根据角色不同将呈现不同功能入口，系统总用例图如图2-1所示。")
    p(doc, "图2-1 系统总用例图")
    p(doc, "（1）用户端角色用例")
    p(
        doc,
        "用户首先在首页与分类页检索商品，随后将商品加入购物车并在结算页提交订单，最终完成支付和订单追踪。其用例图如图2-2所示。"
    )
    p(doc, "图2-2 用户端用例图")
    p(doc, "（2）商家端角色用例")
    p(
        doc,
        "商家首先维护商品与花材信息，随后处理待办订单并监控库存预警，最终根据销售趋势与补货建议完成运营决策。其用例图如图2-3所示。"
    )
    p(doc, "图2-3 商家端用例图")
    p(doc, "（3）管理员端角色用例")
    p(
        doc,
        "管理员首先审核商家与用户数据，随后监管订单处理过程并维护系统参数，最终保障平台运行稳定。其用例图如图2-4所示。"
    )
    p(doc, "图2-4 管理员端用例图")

    p(doc, "2.3 非功能性需求分析")
    p(doc, "除功能性需求外，系统还需满足性能、可维护性与运行环境要求。")

    p(doc, "2.3.1 系统性能需求分析")
    p(
        doc,
        "（1）并发处理需求：在节日高峰场景下，系统应保证订单创建与库存锁定操作具备原子性，避免重复扣减和超卖。"
    )
    p(
        doc,
        "（2）响应时延需求：商品查询、订单详情查询等高频接口在常规负载下应保持秒级响应，保障前端交互流畅。"
    )
    p(
        doc,
        "（3）可靠性需求：订单取消、退款回滚和超时释放必须可追踪、可恢复，确保库存账实一致。"
    )

    p(doc, "2.3.2 开发环境")
    p(doc, "（1）开发语言")
    p(
        doc,
        "后端采用 Java 17 与 Spring Boot 3，前端采用 Vue3 与微信小程序原生框架，分析层采用 Python。"
    )
    p(doc, "（2）开发环境")
    p(
        doc,
        "后端使用 Maven 构建，数据库为 MySQL 8，接口调试与文档通过 Swagger/OpenAPI 支撑，前端使用 Vite 工具链完成构建与发布。"
    )

    p(doc, "3 概要设计")
    p(
        doc,
        "本章在需求分析基础上，完成系统架构、功能模块与数据库三部分概要设计，为详细实现提供统一约束。"
    )

    p(doc, "3.1 系统架构设计")
    p(
        doc,
        "系统采用展示层、业务层、数据层与分析层四层协同架构。用户端与管理端通过 HTTP 调用后端 REST 接口，后端服务负责订单、库存、售后与统计处理，数据落库到 MySQL，分析服务周期生成预测与推荐结果。系统架构如图3-1所示。"
    )
    p(doc, "图3-1 系统架构图")

    p(doc, "3.2 系统功能模块设计")
    p(
        doc,
        "系统按角色划分为用户端、商家端、管理员端三类模块。模块间通过统一接口协作，形成从交易到治理的业务闭环。下面将对各角色功能模块进行说明。"
    )

    p(doc, "3.2.1 用户端功能模块")
    p(doc, "用户端功能模块有商品浏览、购物车、下单支付、订单管理、收货地址管理等功能，功能模块图如图3-2所示。")
    p(doc, "图3-2 用户端功能模块图")
    p(doc, "（1）商品浏览：提供首页推荐、分类筛选与详情展示能力。")
    p(doc, "（2）购物车结算：支持商品增删改查与统一结算提交。")
    p(doc, "（3）下单支付：完成订单创建、支付确认、状态查询与取消操作。")
    p(doc, "（4）地址与评价：支持收货地址维护和订单评价反馈。")

    p(doc, "3.2.2 商家端功能模块")
    p(doc, "商家端功能模块有商品维护、订单处理、库存管理、预警查看、销售报表等功能，功能模块图如图3-3所示。")
    p(doc, "图3-3 商家端功能模块图")
    p(doc, "（1）商品维护：管理商品基础信息、分类、价格与上下架状态。")
    p(doc, "（2）订单处理：查询订单、确认履约、处理异常订单与售后工单。")
    p(doc, "（3）库存管理：执行批次入库、低库存预警和 FEFO 出库管理。")
    p(doc, "（4）经营分析：查看销售趋势、热销商品和补货建议结果。")

    p(doc, "3.2.3 管理员端功能模块")
    p(doc, "管理员端功能模块有用户管理、商家管理、订单监管、系统配置等功能，功能模块图如图3-4所示。")
    p(doc, "图3-4 管理员端功能模块图")
    p(doc, "（1）用户治理：维护用户基础数据并处理异常账号。")
    p(doc, "（2）商家治理：审核商家状态并管理平台商家信息。")
    p(doc, "（3）订单监管：对关键订单状态进行核查并处理治理工单。")
    p(doc, "（4）系统配置：维护全局参数与 Banner 配置，保障平台稳定。")

    p(doc, "3.3 数据库设计")
    p(doc, "数据库设计包含概念结构设计与逻辑结构设计。")

    p(doc, "3.3.1 概念结构设计")
    p(
        doc,
        "在上述功能模块设计结束后，基本可以确定实体，其中包括用户实体、商品实体、订单实体、库存批次实体、售后记录实体。下面将对每个实体进行描述。"
    )
    p(doc, "（1）用户实体：该实体属性包括用户编号、openid、姓名、手机号、积分、创建时间。具体如图3-5所示。")
    p(doc, "图3-5 用户实体属性描述")
    p(doc, "（2）商品实体：该实体属性包括商品编号、标题、分类、基础价格、包装费、配送费、状态。具体如图3-6所示。")
    p(doc, "图3-6 商品实体属性描述")
    p(doc, "（3）订单实体：该实体属性包括订单编号、用户编号、订单状态、总金额、支付金额、收货信息。具体如图3-7所示。")
    p(doc, "图3-7 订单实体属性描述")
    p(doc, "（4）库存批次实体：该实体属性包括批次编号、花材编号、到货时间、枯萎时间、当前库存、锁定库存。具体如图3-8所示。")
    p(doc, "图3-8 库存批次实体属性描述")
    p(doc, "（5）售后记录实体：该实体属性包括售后编号、订单编号、退款金额、申请时间、审核时间、处理状态。具体如图3-9所示。")
    p(doc, "图3-9 售后记录实体属性描述")

    p(doc, "\uff086\uff09\u8ba2\u5355\u660e\u7ec6\u5b9e\u4f53\uff1a\u8be5\u5b9e\u4f53\u5c5e\u6027\u5305\u62ec\u660e\u7ec6\u7f16\u53f7\u3001\u8ba2\u5355\u7f16\u53f7\u3001\u5546\u54c1\u7f16\u53f7\u3001\u5546\u54c1\u540d\u79f0\u3001\u5355\u4ef7\u3001\u6570\u91cf\u3001\u5c0f\u8ba1\u91d1\u989d\u3001\u521b\u5efa\u65f6\u95f4\u3002\u5177\u4f53\u5982\u56fe3-10\u6240\u793a\u3002")
    p(doc, "\u56fe3-10 \u8ba2\u5355\u660e\u7ec6\u5b9e\u4f53\u5c5e\u6027\u63cf\u8ff0")
    p(doc, "\uff087\uff09\u5e93\u5b58\u9501\u6d41\u6c34\u5b9e\u4f53\uff1a\u8be5\u5b9e\u4f53\u5c5e\u6027\u5305\u62ec\u6d41\u6c34\u7f16\u53f7\u3001\u8ba2\u5355\u7f16\u53f7\u3001\u8ba2\u5355\u53f7\u3001\u82b1\u6750\u7f16\u53f7\u3001\u6279\u6b21\u7f16\u53f7\u3001\u9501\u5b9a\u6570\u91cf\u3001\u9501\u5b9a\u72b6\u6001\u3001\u8fc7\u671f\u65f6\u95f4\u3002\u5177\u4f53\u5982\u56fe3-11\u6240\u793a\u3002")
    p(doc, "\u56fe3-11 \u5e93\u5b58\u9501\u6d41\u6c34\u5b9e\u4f53\u5c5e\u6027\u63cf\u8ff0")
    p(doc, "\uff088\uff09\u652f\u4ed8\u6d41\u6c34\u5b9e\u4f53\uff1a\u8be5\u5b9e\u4f53\u5c5e\u6027\u5305\u62ec\u6d41\u6c34\u7f16\u53f7\u3001\u8ba2\u5355\u7f16\u53f7\u3001\u8ba2\u5355\u53f7\u3001\u652f\u4ed8\u4ea4\u6613\u53f7\u3001\u652f\u4ed8\u6e20\u9053\u3001\u652f\u4ed8\u91d1\u989d\u3001\u56de\u8c03\u65f6\u95f4\u3001\u5904\u7406\u7ed3\u679c\u3002\u5177\u4f53\u5982\u56fe3-12\u6240\u793a\u3002")
    p(doc, "\u56fe3-12 \u652f\u4ed8\u6d41\u6c34\u5b9e\u4f53\u5c5e\u6027\u63cf\u8ff0")
    p(doc, "\uff089\uff09\u8bc4\u4ef7\u5b9e\u4f53\uff1a\u8be5\u5b9e\u4f53\u5c5e\u6027\u5305\u62ec\u8bc4\u4ef7\u7f16\u53f7\u3001\u8ba2\u5355\u7f16\u53f7\u3001\u7528\u6237\u7f16\u53f7\u3001\u8bc4\u5206\u3001\u8bc4\u4ef7\u5185\u5bb9\u3001\u8bc4\u4ef7\u6807\u7b7e\u3001\u8bc4\u4ef7\u65f6\u95f4\u3002\u5177\u4f53\u5982\u56fe3-13\u6240\u793a\u3002")
    p(doc, "\u56fe3-13 \u8bc4\u4ef7\u5b9e\u4f53\u5c5e\u6027\u63cf\u8ff0")
    p(doc, "\uff0810\uff09\u9884\u6d4b\u7ed3\u679c\u5b9e\u4f53\uff1a\u8be5\u5b9e\u4f53\u5c5e\u6027\u5305\u62ec\u7ed3\u679c\u7f16\u53f7\u3001\u82b1\u6750\u7f16\u53f7\u3001\u9884\u6d4b\u65e5\u671f\u3001\u9884\u6d4b\u9500\u91cf\u3001\u7f6e\u4fe1\u4e0b\u9650\u3001\u7f6e\u4fe1\u4e0a\u9650\u3001\u6a21\u578b\u540d\u79f0\u3001\u751f\u6210\u65f6\u95f4\u3002\u5177\u4f53\u5982\u56fe3-14\u6240\u793a\u3002")
    p(doc, "\u56fe3-14 \u9884\u6d4b\u7ed3\u679c\u5b9e\u4f53\u5c5e\u6027\u63cf\u8ff0")
    p(doc, "\uff0811\uff09\u8865\u8d27\u5efa\u8bae\u5b9e\u4f53\uff1a\u8be5\u5b9e\u4f53\u5c5e\u6027\u5305\u62ec\u5efa\u8bae\u7f16\u53f7\u3001\u82b1\u6750\u7f16\u53f7\u3001\u5efa\u8bae\u65e5\u671f\u3001\u9884\u6d4b\u9700\u6c42\u3001\u5b89\u5168\u5e93\u5b58\u3001\u518d\u8ba2\u8d27\u70b9\u3001\u73b0\u6709\u5e93\u5b58\u3001\u5728\u9014\u6570\u91cf\u3001\u5efa\u8bae\u8865\u8d27\u91cf\u3001\u5efa\u8bae\u72b6\u6001\u3002\u5177\u4f53\u5982\u56fe3-15\u6240\u793a\u3002")
    p(doc, "\u56fe3-15 \u8865\u8d27\u5efa\u8bae\u5b9e\u4f53\u5c5e\u6027\u63cf\u8ff0")
    p(doc, "\uff0812\uff09\u63a8\u8350\u7ed3\u679c\u5b9e\u4f53\uff1a\u8be5\u5b9e\u4f53\u5c5e\u6027\u5305\u62ec\u63a8\u8350\u7f16\u53f7\u3001\u7528\u6237\u7f16\u53f7\u3001\u5546\u54c1\u7f16\u53f7\u3001\u63a8\u8350\u5f97\u5206\u3001\u63a8\u8350\u539f\u56e0\u3001\u751f\u6210\u65f6\u95f4\u3002\u5177\u4f53\u5982\u56fe3-16\u6240\u793a\u3002")
    p(doc, "\u56fe3-16 \u63a8\u8350\u7ed3\u679c\u5b9e\u4f53\u5c5e\u6027\u63cf\u8ff0")

    p(doc, "\uff0813\uff09\u5546\u5bb6\u5b9e\u4f53\uff1a\u8be5\u5b9e\u4f53\u5c5e\u6027\u5305\u62ec\u5546\u5bb6\u7f16\u53f7\u3001\u5546\u5bb6\u540d\u79f0\u3001\u8054\u7cfb\u7535\u8bdd\u3001\u90ae\u7bb1\u3001\u5730\u5740\u3001\u72b6\u6001\u3001\u521b\u5efa\u65f6\u95f4\u3001\u66f4\u65b0\u65f6\u95f4\u3002\u5177\u4f53\u5982\u56fe3-17\u6240\u793a\u3002")
    p(doc, "\u56fe3-17 \u5546\u5bb6\u5b9e\u4f53\u5c5e\u6027\u63cf\u8ff0")
    p(doc, "\uff0814\uff09\u5206\u7c7b\u5b9e\u4f53\uff1a\u8be5\u5b9e\u4f53\u5c5e\u6027\u5305\u62ec\u5206\u7c7b\u7f16\u53f7\u3001\u5206\u7c7b\u7f16\u7801\u3001\u5206\u7c7b\u540d\u79f0\u3001\u6392\u5e8f\u3001\u56fe\u6807\u3001\u542f\u7528\u72b6\u6001\u3001\u521b\u5efa\u65f6\u95f4\u3001\u66f4\u65b0\u65f6\u95f4\u3002\u5177\u4f53\u5982\u56fe3-18\u6240\u793a\u3002")
    p(doc, "\u56fe3-18 \u5206\u7c7b\u5b9e\u4f53\u5c5e\u6027\u63cf\u8ff0")
    p(doc, "\uff0815\uff09\u8d2d\u7269\u8f66\u9879\u5b9e\u4f53\uff1a\u8be5\u5b9e\u4f53\u5c5e\u6027\u5305\u62ec\u8d2d\u7269\u8f66\u9879\u7f16\u53f7\u3001\u7528\u6237\u7f16\u53f7\u3001\u5546\u54c1\u7f16\u53f7\u3001\u6570\u91cf\u3001\u521b\u5efa\u65f6\u95f4\u3001\u66f4\u65b0\u65f6\u95f4\u3002\u5177\u4f53\u5982\u56fe3-19\u6240\u793a\u3002")
    p(doc, "\u56fe3-19 \u8d2d\u7269\u8f66\u9879\u5b9e\u4f53\u5c5e\u6027\u63cf\u8ff0")
    p(doc, "\uff0816\uff09\u7cfb\u7edf\u914d\u7f6e\u5b9e\u4f53\uff1a\u8be5\u5b9e\u4f53\u5c5e\u6027\u5305\u62ec\u914d\u7f6e\u7f16\u53f7\u3001\u914d\u7f6e\u952e\u3001\u914d\u7f6e\u503c\u3001\u63cf\u8ff0\u3001\u66f4\u65b0\u65f6\u95f4\u3002\u5177\u4f53\u5982\u56fe3-20\u6240\u793a\u3002")
    p(doc, "\u56fe3-20 \u7cfb\u7edf\u914d\u7f6e\u5b9e\u4f53\u5c5e\u6027\u63cf\u8ff0")
    p(doc, "\uff0817\uff09\u8f6e\u64ad\u56fe\u5b9e\u4f53\uff1a\u8be5\u5b9e\u4f53\u5c5e\u6027\u5305\u62ec\u8f6e\u64ad\u56fe\u7f16\u53f7\u3001\u6807\u9898\u3001\u526f\u6807\u9898\u3001\u8d77\u59cb\u989c\u8272\u3001\u7ed3\u675f\u989c\u8272\u3001\u8df3\u8f6c\u94fe\u63a5\u3001\u6392\u5e8f\u3001\u542f\u7528\u72b6\u6001\u3001\u521b\u5efa\u65f6\u95f4\u3001\u66f4\u65b0\u65f6\u95f4\u3002\u5177\u4f53\u5982\u56fe3-21\u6240\u793a\u3002")
    p(doc, "\u56fe3-21 \u8f6e\u64ad\u56fe\u5b9e\u4f53\u5c5e\u6027\u63cf\u8ff0")

    p(doc, "3.3.2 逻辑结构设计")
    p(
        doc,
        "数据库表结构的逻辑设计可以直白地体现出每个字段属性的含义，下面将对每个表的逻辑结构设计进行叙述。首先是对数据库结构进行描述，数据库结构如表3-1所示。"
    )
    table_with_caption(
        doc,
        "表3-1 数据库结构",
        ["序号", "表名", "说明"],
        [
            ["1", "user_customer", "用户基础信息表"],
            ["2", "product", "商品信息表"],
            ["3", "customer_order", "订单主表"],
            ["4", "order_item", "订单明细表"],
            ["5", "inventory_batch", "库存批次表"],
            ["6", "stock_lock", "库存锁流水表"],
            ["7", "after_sale_record", "售后记录表"],
            ["8", "forecast_result", "销量预测结果表"],
            ["9", "replenishment_suggestion", "补货建议表"],
        ],
    )

    p(doc, "(1)用户表（user_customer）")
    p(doc, "该表用于记录系统用户基础信息，其属性包括 id、openid、name、phone、points、created_at、updated_at，用户表的表结构如表3-2所示。")
    table_with_caption(
        doc,
        "表3-2 用户表表结构",
        ["字段名", "描述", "数据类型", "是否为空", "约束"],
        [
            ["id", "用户主键", "BIGINT", "否", "PK, AUTO_INCREMENT"],
            ["openid", "微信标识", "VARCHAR(64)", "否", "UNIQUE"],
            ["name", "用户姓名", "VARCHAR(64)", "否", "-"],
            ["phone", "手机号", "VARCHAR(20)", "是", "UNIQUE"],
            ["points", "积分", "INT", "否", "DEFAULT 0"],
        ],
    )

    p(doc, "(2)商品表（product）")
    p(doc, "该表用于记录商品展示与价格信息，其属性包括 id、title、type、category、base_price、packaging_fee、delivery_fee、status，商品表的表结构如表3-3所示。")
    table_with_caption(
        doc,
        "表3-3 商品表表结构",
        ["字段名", "描述", "数据类型", "是否为空", "约束"],
        [
            ["id", "商品主键", "BIGINT", "否", "PK, AUTO_INCREMENT"],
            ["title", "商品标题", "VARCHAR(128)", "否", "-"],
            ["type", "商品类型", "VARCHAR(16)", "否", "-"],
            ["category", "商品分类", "VARCHAR(32)", "否", "-"],
            ["status", "销售状态", "VARCHAR(16)", "否", "DEFAULT ON_SALE"],
        ],
    )

    p(doc, "(3)订单表（customer_order）")
    p(doc, "该表用于记录用户订单全生命周期信息，其属性包括 id、order_no、user_id、total_amount、payment_amount、status、pay_time、cancel_time，订单表的表结构如表3-4所示。")
    table_with_caption(
        doc,
        "表3-4 订单表表结构",
        ["字段名", "描述", "数据类型", "是否为空", "约束"],
        [
            ["id", "订单主键", "BIGINT", "否", "PK, AUTO_INCREMENT"],
            ["order_no", "订单编号", "VARCHAR(40)", "否", "UNIQUE"],
            ["user_id", "用户编号", "BIGINT", "否", "FK -> user_customer.id"],
            ["total_amount", "订单总额", "DECIMAL(12,2)", "否", "DEFAULT 0.00"],
            ["status", "订单状态", "VARCHAR(20)", "否", "-"],
        ],
    )

    p(doc, "(4)库存锁流水表（stock_lock）")
    p(doc, "该表用于记录订单锁库存与回滚过程信息，其属性包括 order_id、order_no、flower_id、batch_id、lock_qty、status、expires_at，库存锁流水表的表结构如表3-5所示。")
    table_with_caption(
        doc,
        "表3-5 库存锁流水表表结构",
        ["字段名", "描述", "数据类型", "是否为空", "约束"],
        [
            ["id", "流水主键", "BIGINT", "否", "PK, AUTO_INCREMENT"],
            ["order_id", "订单编号", "BIGINT", "否", "FK -> customer_order.id"],
            ["batch_id", "批次编号", "BIGINT", "否", "FK -> inventory_batch.id"],
            ["lock_qty", "锁定数量", "DECIMAL(10,2)", "否", "-"],
            ["status", "锁状态", "VARCHAR(16)", "否", "LOCKED/CONFIRMED/RELEASED/ROLLED_BACK"],
        ],
    )

    p(doc, "(5)售后记录表（after_sale_record）")
    p(doc, "该表用于记录退款申请、审核与处理结果，其属性包括 refund_no、order_no、refund_amount、status、apply_time、audit_time、refund_time，售后记录表的表结构如表3-6所示。")
    table_with_caption(
        doc,
        "表3-6 售后记录表表结构",
        ["字段名", "描述", "数据类型", "是否为空", "约束"],
        [
            ["id", "售后主键", "BIGINT", "否", "PK, AUTO_INCREMENT"],
            ["refund_no", "退款单号", "VARCHAR(32)", "否", "UNIQUE"],
            ["order_no", "订单编号", "VARCHAR(40)", "否", "-"],
            ["refund_amount", "退款金额", "DECIMAL(12,2)", "否", "-"],
            ["status", "处理状态", "VARCHAR(20)", "否", "-"],
        ],
    )

    p(doc, "4 详细设计与实现")
    p(doc, "本章围绕核心业务流程，先给出系统功能设计，再说明对应功能实现。")

    p(doc, "4.1 系统功能设计")
    p(
        doc,
        "系统功能设计围绕订单、库存、售后、分析和治理五条关键链路展开，通过流程化组织实现多角色协同。下面将会对上述功能进行讲解。"
    )

    p(doc, "4.1.1 下单与库存锁定流程设计")
    p(
        doc,
        "首先用户在购物车中确认商品与配送信息并提交订单；随后后端根据商品 BOM 计算花材需求并按 FEFO 规则锁定可用批次库存；最终生成状态为 LOCKED 的订单与锁流水，形成可追踪的下单链路。具体流程如图4-1所示。"
    )
    p(doc, "图4-1 下单与库存锁定流程图")

    p(doc, "4.1.2 支付确认与库存扣减流程设计")
    p(
        doc,
        "首先用户发起支付并提交支付流水号；随后系统校验订单状态和金额后将锁库存转为确认扣减；最终订单状态由 LOCKED 进入 PAID/CONFIRMED，保证库存台账与支付结果一致。具体流程如图4-2所示。"
    )
    p(doc, "图4-2 支付确认与库存扣减流程图")

    p(doc, "4.1.3 取消/退款回滚流程设计")
    p(
        doc,
        "首先系统识别未支付取消或已支付退款两类场景；随后分别执行释放锁库存或回补已扣减库存操作；最终订单状态进入 CANCELLED 或 REFUNDED，并写入售后记录，形成闭环回滚流程。具体流程如图4-3所示。"
    )
    p(doc, "图4-3 取消退款回滚流程图")

    p(doc, "4.1.4 智能推荐与补货流程设计")
    p(
        doc,
        "首先分析任务读取历史订单和库存数据；随后计算销量预测、推荐分值与补货建议；最终将分析结果回写至结果表并在商家端展示，形成“交易驱动分析、分析反哺运营”的数据流程。具体流程如图4-4所示。"
    )
    p(doc, "图4-4 推荐与补货联动流程图")

    p(doc, "4.1.5 商家运营治理流程设计")
    p(
        doc,
        "首先商家在管理端查看待处理订单、库存预警和销售看板；随后按优先级执行商品维护、订单处理与补货计划；最终通过管理员监管与系统参数配置完成治理闭环。具体流程如图4-5所示。"
    )
    p(doc, "图4-5 商家运营治理流程图")

    p(doc, "4.2 系统功能实现")
    p(doc, "本节按照与 4.1 一致的顺序，给出关键功能的实现证据与处理逻辑。")

    p(doc, "4.2.1 下单与库存锁定功能实现")
    p(
        doc,
        "前端入口位于小程序结算页 pages/checkout/checkout.js，提交 POST /api/v1/orders 请求；后端由 OrderController 的 createOrder 方法接收并调用 OrderService 完成订单创建与库存锁定；处理过程中关联 customer_order、order_item、stock_lock、inventory_batch 表，返回订单号、锁定状态和过期时间。运行效果图如图4-6所示。"
    )
    p(doc, "图4-6 下单与库存锁定页面效果图")

    p(doc, "4.2.2 支付确认与库存扣减功能实现")
    p(
        doc,
        "前端入口位于小程序支付页 pages/pay/pay.js，提交 POST /api/v1/orders/{orderNo}/pay 请求；后端由 OrderController 的 payOrder 方法及 PaymentController 回调接口协同处理，调用 PaymentService 更新支付流水并确认库存扣减；核心涉及 customer_order、stock_lock、payment_log 表。运行效果图如图4-7所示。"
    )
    p(doc, "图4-7 支付确认与库存扣减页面效果图")

    p(doc, "4.2.3 取消/退款回滚功能实现")
    p(
        doc,
        "前端入口位于订单详情页 pages/order-detail/order-detail.js，提交 POST /api/v1/orders/{orderNo}/cancel 与 POST /api/v1/after-sales 请求；后端由 OrderController 和 AfterSaleController 联动调用 OrderService、AfterSaleService，分别执行锁库存释放与退款回补；相关数据写入 customer_order、stock_lock、after_sale_record。运行效果图如图4-8所示。"
    )
    p(doc, "图4-8 取消退款回滚页面效果图")

    p(doc, "4.2.4 智能推荐与补货功能实现")
    p(
        doc,
        "商家端入口为 /merchant/sales-report 页面，查询 GET /api/v1/analysis/replenishment 与 GET /api/v1/analysis/recommendations；分析服务通过 forecast_job.py 与 recommendation_job.py 定时计算并回写 forecast_result、replenishment_suggestion、recommendation_result；后端 AnalysisController 完成查询封装。运行效果图如图4-9所示。"
    )
    p(doc, "图4-9 智能推荐与补货页面效果图")

    p(doc, "4.2.5 商家运营治理功能实现")
    p(
        doc,
        "商家与管理员分别从 /merchant/orders、/merchant/inventory-alert、/admin/users、/admin/merchants、/admin/orders 入口开展治理；后端由 StatsController、InventoryController、MerchantController、UserController 提供支撑接口，结合 system_config 与 banner 参数实现运营面板可配置化。运行效果图如图4-10所示。"
    )
    p(doc, "图4-10 商家运营治理页面效果图")

    p(doc, "5 系统测试")
    p(
        doc,
        "系统测试采用黑盒测试中的等价类划分法，围绕订单、库存、售后与分析四个核心能力进行验证，重点确认业务正确性与状态一致性。"
    )

    p(doc, "5.1 下单锁库存功能测试")
    p(
        doc,
        "本节验证用户提交订单后库存是否按 FEFO 正确锁定，覆盖正常下单、库存不足和并发下单等等价类场景，其测试用例如表5-1所示。"
    )
    table_with_caption(
        doc,
        "表5-1 下单锁库存功能测试用例",
        ["编号", "测试内容", "执行操作", "预期结果", "测试结果"],
        [
            ["TC5-1-01", "正常下单", "提交单商品订单", "返回 LOCKED 且生成锁流水", "通过"],
            ["TC5-1-02", "库存不足", "提交超库存订单", "返回失败并提示库存不足", "通过"],
            ["TC5-1-03", "并发下单", "并发提交同商品订单", "无超卖，锁量总和不超过可用量", "通过"],
        ],
    )

    p(doc, "5.2 支付扣减与库存一致性测试")
    p(
        doc,
        "本节验证支付成功后库存是否从锁定态转为确认扣减，并校验支付日志记录完整性，其测试用例如表5-2所示。"
    )
    table_with_caption(
        doc,
        "表5-2 支付扣减与库存一致性测试用例",
        ["编号", "测试内容", "执行操作", "预期结果", "测试结果"],
        [
            ["TC5-2-01", "支付成功", "调用支付确认接口", "订单状态更新为 PAID，锁库存转 CONFIRMED", "通过"],
            ["TC5-2-02", "重复支付", "对同订单重复调用支付", "系统幂等处理，不重复扣减库存", "通过"],
            ["TC5-2-03", "支付日志", "查询 payment_log", "存在对应交易流水记录", "通过"],
        ],
    )

    p(doc, "5.3 退款回滚与售后流程测试")
    p(
        doc,
        "本节验证取消订单与售后退款场景下库存回滚逻辑是否正确，覆盖取消未支付、审核退款和异常拒绝三类场景，其测试用例如表5-3所示。"
    )
    table_with_caption(
        doc,
        "表5-3 退款回滚与售后流程测试用例",
        ["编号", "测试内容", "执行操作", "预期结果", "测试结果"],
        [
            ["TC5-3-01", "取消未支付订单", "调用取消接口", "订单状态变更为 CANCELLED，锁库存释放", "通过"],
            ["TC5-3-02", "退款审核通过", "申请并处理退款", "订单状态为 REFUNDED，库存完成回补", "通过"],
            ["TC5-3-03", "退款审核拒绝", "执行拒绝操作", "售后记录状态为 REJECTED，不触发库存回补", "通过"],
        ],
    )

    p(doc, "5.4 推荐与补货结果可用性测试")
    p(
        doc,
        "本节验证分析任务是否可稳定生成推荐与补货结果，覆盖任务执行、结果查询和异常数据兜底三类场景，其测试用例如表5-4所示。"
    )
    table_with_caption(
        doc,
        "表5-4 推荐与补货结果可用性测试用例",
        ["编号", "测试内容", "执行操作", "预期结果", "测试结果"],
        [
            ["TC5-4-01", "预测任务执行", "触发 forecast 任务", "生成未来销量预测结果", "通过"],
            ["TC5-4-02", "补货建议查询", "查询 replenishment 接口", "返回建议补货量与安全库存", "通过"],
            ["TC5-4-03", "推荐结果查询", "查询 recommendations 接口", "返回推荐商品与推荐理由", "通过"],
        ],
    )
    p(
        doc,
        "测试结果表明，系统核心链路能够稳定运行，订单状态、库存状态与售后状态之间保持一致，分析模块输出结果可用于商家运营决策。"
    )

    p(doc, "6 总结")
    p(
        doc,
        "本文围绕花店线上经营场景，完成了基于微信小程序的花店智能管理系统设计与实现。系统在交易流程、库存一致性、售后回滚与运营分析四个方面形成可运行的闭环能力，并通过分层架构保障模块扩展性。"
    )
    p(
        doc,
        "工程实践表明，FEFO 批次策略与库存锁机制能够有效降低超卖与错卖风险，推荐与补货结果可支撑商家日常决策。后续工作可在真实生产数据规模下进一步优化预测模型精度、引入更细粒度权限审计与自动化告警策略。"
    )

    p(doc, "参考文献")
    refs = [
        "[1] 王珊, 萨师煊. 数据库系统概论(第6版)[M]. 北京: 高等教育出版社, 2023.",
        "[2] 刘文全. 在线花店管理系统设计与实现[J]. 现代商贸工业, 2009.",
        "[3] 杨柯. 基于B/S模式的鲜花销售管理系统设计[J]. 民营科技, 2016.",
        "[4] 段阳. 基于B/S结构的网上花店管理系统设计研究[J]. 电脑知识与技术, 2019.",
        "[5] 微信开放社区. 微信小程序开发文档[EB/OL]. https://developers.weixin.qq.com/miniprogram/dev/framework/.",
        "[6] 中国信息通信研究院. 中国数字经济发展研究报告[R]. 北京: 中国信息通信研究院, 2024.",
        "[7] Craig Walls. Spring in Action, 6th Edition[M]. Manning Publications, 2022.",
        "[8] Oracle. MySQL 8.0 Reference Manual[EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/.",
        "[9] Wes McKinney. Python for Data Analysis, 3rd Edition[M]. O'Reilly Media, 2022.",
        "[10] Rob J Hyndman, George Athanasopoulos. Forecasting: Principles and Practice, 3rd Edition[EB/OL]. https://otexts.com/fpp3/.",
        "[11] Martin Kleppmann. Designing Data-Intensive Applications[M]. O'Reilly Media, 2017.",
        "[12] Kunaver M, Pozrl T. Diversity in recommender systems: A survey[J]. Knowledge-Based Systems, 2017, 123:154-162.",
        "[13] Wang X, He X, Wang M, et al. Neural graph collaborative filtering[C]. SIGIR, 2019.",
        "[14] Hochreiter S, Schmidhuber J. Long short-term memory[J]. Neural Computation, 1997, 9(8):1735-1780.",
        "[15] Vaswani A, Shazeer N, Parmar N, et al. Attention is all you need[C]. NIPS, 2017.",
        "[16] Hodson T O. RMSE or MAE: When to use them or not[J]. Geoscientific Model Development, 2022, 15(14):5481-5487.",
        "[17] 陈强. 软件测试方法与实践[M]. 北京: 机械工业出版社, 2021.",
        "[18] 李航. 统计学习方法(第2版)[M]. 北京: 清华大学出版社, 2019.",
    ]
    for r in refs:
        p(doc, r)

    p(doc, "致谢")
    p(
        doc,
        "本课题完成过程中，首先感谢指导教师在选题方向、系统设计与论文写作方面给予的持续指导，使我能够在工程实现和学术表达上不断完善。"
    )
    p(
        doc,
        "感谢家人在学习期间给予的理解与支持，感谢项目联调与测试阶段给予帮助的同学与朋友。"
    )
    p(
        doc,
        "同时感谢相关研究文献作者与开源社区提供的技术资料，为本课题的实现与优化提供了重要参考。"
    )

    doc.save(output_path)


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[2]
    out = root / "花店智能管理系统-毕业论文-初稿.docx"
    build_doc(out)
    print(f"generated: {out}")
