from __future__ import annotations

import argparse
import re
from collections import Counter
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def parse_audit_count(text: str, key: str) -> int:
    m = re.search(rf"-\s*{re.escape(key)}:\s*(\d+)", text)
    return int(m.group(1)) if m else -1


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--docx", required=True)
    parser.add_argument("--audit", required=True)
    parser.add_argument("--workdir", required=True)
    args = parser.parse_args()

    docx_path = Path(args.docx)
    audit_path = Path(args.audit)
    work = Path(args.workdir)

    d = Document(docx_path)
    paras = d.paragraphs

    # 1) English leak check
    english_leak_hits = []
    in_en_abstract = False
    in_references = False
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not t:
            continue
        if t == "参考文献":
            in_references = True
            continue
        if t == "致谢":
            in_references = False
        if re.match(r"^Abstract\b", t, flags=re.IGNORECASE):
            in_en_abstract = True
            continue
        if re.match(r"^Key\s*words\b", t, flags=re.IGNORECASE):
            in_en_abstract = False
            continue
        if in_en_abstract:
            continue
        if in_references and re.match(r"^\[\d+\]", t):
            continue
        compact = re.sub(r"\s+", "", t)
        if len(compact) < 120:
            continue
        letters = len(re.findall(r"[A-Za-z]", t))
        cjk = len(re.findall(r"[\u4e00-\u9fff]", t))
        if letters >= 100 and cjk <= 20 and letters > cjk * 3:
            english_leak_hits.append((i + 1, t[:160]))

    # 2) Placeholder check
    placeholder_hits = []
    for i, p in enumerate(paras):
        t = p.text or ""
        if re.search(r"\?{2,}", t):
            placeholder_hits.append((i + 1, t[:160]))

    # 2.1) broken-word / mojibake check
    broken_word_hits = []
    for i, p in enumerate(paras):
        t = p.text or ""
        if re.search(r"\?{3,}", t) or re.search(r"[A-Za-z]\?[A-Za-z]", t):
            broken_word_hits.append((i + 1, t[:160]))

    # 2.2) table-cell placeholder check
    table_cell_placeholder_hits = []
    for ti, t in enumerate(d.tables):
        for r, row in enumerate(t.rows):
            for c, cell in enumerate(row.cells):
                txt = cell.text or ""
                if re.search(r"\?{2,}", txt):
                    table_cell_placeholder_hits.append((ti + 1, r + 1, c + 1, txt[:120]))

    # 2.3) draft placeholder check
    draft_placeholder_hits = []
    for i, p in enumerate(paras):
        t = p.text or ""
        if re.search(r"(?<![A-Za-z])x{3,}(?![A-Za-z])", t, flags=re.IGNORECASE):
            draft_placeholder_hits.append((i + 1, "xxx", t[:160]))
        elif re.search(r"【待补[:：][^】]*】", t):
            draft_placeholder_hits.append((i + 1, "todo_cn", t[:160]))

    # 3) body heading-prefix check
    body_heading_prefix_hits = []
    pat = re.compile(r"^\s*\d+(?:\.\d+){0,2}\s+\S+")
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not t or not pat.match(t):
            continue
        style_name = (p.style.name if p.style else "").strip()
        style_id = (p.style.style_id if p.style else "").strip().lower()
        is_heading = bool(re.match(r"^(\u6807\u9898|Heading)\s*[1-3]$", style_name, flags=re.IGNORECASE))
        is_heading = is_heading or style_id in ("heading1", "heading2", "heading3")
        if not is_heading:
            body_heading_prefix_hits.append((i + 1, style_name, style_id, t[:120]))

    # 4) heading style check
    heading_style_mismatch_hits = []
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not pat.match(t):
            continue
        level = min(t.split()[0].count(".") + 1, 3)
        style_name = (p.style.name if p.style else "").strip()
        style_id = (p.style.style_id if p.style else "").strip().lower()
        name_ok = bool(re.match(rf"^(\u6807\u9898|Heading)\s*{level}$", style_name, flags=re.IGNORECASE))
        id_ok = style_id == f"heading{level}"
        if not (name_ok or id_ok):
            heading_style_mismatch_hits.append((i + 1, level, style_name, style_id, t[:120]))

    # 5) full-template conformance check (project_analysis)
    seen = []
    num_title_pat = re.compile(r"^\s*(\d+(?:\.\d+){0,2})\s+(.+?)\s*$")
    for p in paras:
        t = (p.text or "").strip()
        m = num_title_pat.match(t)
        if m:
            seen.append((m.group(1), m.group(2)))
    seen_nums = [n for n, _ in seen]
    seen_title_map = {}
    for n, title in seen:
        if n not in seen_title_map:
            seen_title_map[n] = title

    required_nums = [
        "1", "1.1", "1.2", "1.3", "1.4",
        "2", "2.1", "2.1.1", "2.1.2", "2.2", "2.2.1", "2.2.2", "2.3", "2.3.1", "2.3.2",
        "3", "3.1", "3.2", "3.3", "3.3.1", "3.3.2",
        "4", "4.1", "4.2",
        "5", "5.1", "5.2", "5.3", "5.4",
        "6",
    ]
    required_exact = {
        "1": "概述",
        "1.1": "课题研究背景和意义",
        "1.2": "国内外研究现状",
        "1.3": "课题主要研究内容",
        "1.4": "论文组织结构",
        "2": "需求分析",
        "2.1": "可行性分析",
        "2.1.1": "经济可行性分析",
        "2.1.2": "技术可行性分析",
        "2.2": "功能性需求分析",
        "2.2.1": "系统功能概述",
        "2.2.2": "系统功能需求分析",
        "2.3": "非功能性需求分析",
        "2.3.1": "系统性能需求分析",
        "2.3.2": "开发环境",
        "3": "概要设计",
        "3.1": "系统架构设计",
        "3.2": "系统功能模块设计",
        "3.3": "数据库设计",
        "3.3.1": "概念结构设计",
        "3.3.2": "逻辑结构设计",
        "4": "详细设计与实现",
        "4.1": "系统功能设计",
        "4.2": "系统功能实现",
        "5": "系统测试",
        "6": "总结",
    }
    missing_nums = [n for n in required_nums if n not in seen_nums]
    dup_nums = sorted([n for n, c in Counter(seen_nums).items() if c > 1 and n in required_nums])
    exact_mismatch = []
    for n, expected in required_exact.items():
        if n not in seen_title_map:
            continue
        actual = re.sub(r"\s+", "", seen_title_map[n])
        want = re.sub(r"\s+", "", expected)
        if actual != want:
            exact_mismatch.append((n, seen_title_map[n], expected))
    old_ch5_titles = {
        "5.1": "课堂创建功能测试",
        "5.2": "上传视频功能测试",
        "5.3": "排行榜查看功能测试",
        "5.4": "在线编译功能测试",
    }
    project_mode_sample_ch5_hits = []
    for n, old_title in old_ch5_titles.items():
        if re.sub(r"\s+", "", seen_title_map.get(n, "")) == re.sub(r"\s+", "", old_title):
            project_mode_sample_ch5_hits.append(f"{n} {old_title}")

    # 6) project-domain leakage check
    forbidden_terms = [
        "在线编译",
        "课堂管理",
        "课程管理",
        "教师功能模块",
        "学生功能模块",
        "课堂创建功能测试",
        "上传视频功能测试",
        "排行榜查看功能测试",
    ]
    project_domain_leak_hits = []
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if not t:
            continue
        for term in forbidden_terms:
            if term in t:
                project_domain_leak_hits.append((i + 1, term, t[:160]))

    # 7) caption completeness
    fig_refs, tab_refs, fig_caps, tab_caps = set(), set(), set(), set()
    for p in paras:
        t = (p.text or "").strip()
        if not t:
            continue
        m = re.match(r"^图(\d+-\d+)\s+", t)
        if m:
            fig_caps.add(m.group(1))
        m = re.match(r"^表(\d+-\d+)\s+", t)
        if m:
            tab_caps.add(m.group(1))
        for x in re.findall(r"(?:如图|见图|图)(\d+-\d+)", t):
            fig_refs.add(x)
        for x in re.findall(r"(?:如表|见表|表)(\d+-\d+)", t):
            tab_refs.add(x)
    missing_fig_caps = sorted(fig_refs - fig_caps)
    missing_tab_caps = sorted(tab_refs - tab_caps)

    # 7.1) diagram generation coverage
    audit_text = audit_path.read_text(encoding="utf-8", errors="replace")
    fig_caption = parse_audit_count(audit_text, "fig_caption")
    figures_inserted = parse_audit_count(audit_text, "figures_inserted")
    figures_pack_exported = parse_audit_count(audit_text, "figures_pack_exported")
    generated = max(figures_inserted, figures_pack_exported)
    diagram_generation_gate_fail = int(fig_caption >= 1 and generated <= 0)

    # 7.2) ordered figure-pack check
    src = work / "图包_手动粘贴_按图序"
    dst = work / "图包_手动粘贴_按图序_编号版"
    src_bmp = sorted(src.glob("*.bmp")) if src.exists() else []
    dst_bmp = sorted(dst.glob("*.bmp")) if dst.exists() else []
    src_map = (src / "图注与图片对照.txt").exists()
    dst_map = (dst / "图注与图片对照.txt").exists()
    seq_ok = True
    for i, pth in enumerate(dst_bmp, start=1):
        if not re.match(rf"^{i:02d}_", pth.name):
            seq_ok = False
            break
    ordered_pack_gate_fail = int(
        not (
            src.exists()
            and dst.exists()
            and len(src_bmp) >= 1
            and len(dst_bmp) >= 1
            and src_map
            and dst_map
            and seq_ok
        )
    )

    # 8) chapter-5 fake inline rows
    chapter5_fake_inline_rows = []
    in_ch5 = False
    for i, p in enumerate(paras):
        t = (p.text or "").strip()
        if re.match(r"^5(\s|\.0*\s)", t):
            in_ch5 = True
        if re.match(r"^6\s+\S+", t):
            in_ch5 = False
        if not in_ch5 or not t:
            continue
        if re.match(r"^\d+\s+", t) and "符合预期" in t:
            chapter5_fake_inline_rows.append((i + 1, t[:160]))

    # 9) three-line table border check
    three_line_table_violations = []
    for i, t in enumerate(d.tables):
        tbl_pr = t._tbl.tblPr
        b = tbl_pr.find(qn("w:tblBorders")) if tbl_pr is not None else None
        if b is None:
            three_line_table_violations.append((i + 1, "missing_tblBorders"))
            continue

        def val(edge: str) -> str | None:
            e = b.find(qn(f"w:{edge}"))
            return e.get(qn("w:val")) if e is not None else None

        top = val("top")
        bottom = val("bottom")
        left = val("left")
        right = val("right")
        inside_v = val("insideV")
        if top in (None, "nil") or bottom in (None, "nil"):
            three_line_table_violations.append((i + 1, "top_or_bottom_missing"))
        if left not in ("nil", "none") or right not in ("nil", "none") or inside_v not in ("nil", "none"):
            three_line_table_violations.append((i + 1, "vertical_border_not_nil"))

    # 9.1) strict three-line table visual check
    tables_with_tblStyle = 0
    cell_vertical_override = []
    cell_three_line_shape = []
    for ti, t in enumerate(d.tables, 1):
        tbl_pr = t._tbl.tblPr
        if tbl_pr is not None and tbl_pr.find(qn("w:tblStyle")) is not None:
            tables_with_tblStyle += 1
        rows = t.rows
        if not rows:
            continue

        def get_edge(tc_borders, edge: str):
            if tc_borders is None:
                return (None, None)
            e = tc_borders.find(qn(f"w:{edge}"))
            if e is None:
                return (None, None)
            return (e.get(qn("w:val")), e.get(qn("w:sz")))

        row_count = len(rows)
        for r_idx, row in enumerate(rows):
            for c_idx, cell in enumerate(row.cells):
                tc_pr = cell._tc.tcPr
                tc_borders = tc_pr.find(qn("w:tcBorders")) if tc_pr is not None else None
                left, _ = get_edge(tc_borders, "left")
                right, _ = get_edge(tc_borders, "right")
                inside_v, _ = get_edge(tc_borders, "insideV")
                inside_h, _ = get_edge(tc_borders, "insideH")
                top, top_sz = get_edge(tc_borders, "top")
                bottom, bottom_sz = get_edge(tc_borders, "bottom")

                if (
                    left not in (None, "nil", "none")
                    or right not in (None, "nil", "none")
                    or inside_v not in (None, "nil", "none")
                    or inside_h not in (None, "nil", "none")
                ):
                    cell_vertical_override.append((ti, r_idx + 1, c_idx + 1))

                if row_count == 1:
                    ok = (top == "single" and top_sz == "12" and bottom == "single" and bottom_sz == "12")
                    if not ok:
                        cell_three_line_shape.append((ti, r_idx + 1, c_idx + 1, "single_row_12_12"))
                elif r_idx == 0:
                    ok = (top == "single" and top_sz == "12" and bottom == "single" and bottom_sz == "4")
                    if not ok:
                        cell_three_line_shape.append((ti, r_idx + 1, c_idx + 1, "header_12_4"))
                elif r_idx == row_count - 1:
                    if not (bottom == "single" and bottom_sz == "12"):
                        cell_three_line_shape.append((ti, r_idx + 1, c_idx + 1, "last_bottom_12"))
                    if top not in (None, "nil", "none"):
                        cell_three_line_shape.append((ti, r_idx + 1, c_idx + 1, "last_top_nil"))
                else:
                    if top not in (None, "nil", "none") or bottom not in (None, "nil", "none"):
                        cell_three_line_shape.append((ti, r_idx + 1, c_idx + 1, "body_nil_nil"))

    # Output summary metrics
    print("english_leak_hits", len(english_leak_hits))
    print("placeholder_hits", len(placeholder_hits))
    print("broken_word_hits", len(broken_word_hits))
    print("table_cell_placeholder_hits", len(table_cell_placeholder_hits))
    print("draft_placeholder_hits", len(draft_placeholder_hits))
    print("body_heading_prefix_hits", len(body_heading_prefix_hits))
    print("heading_style_mismatch_hits", len(heading_style_mismatch_hits))
    print("template_missing_number_count", len(missing_nums))
    print("template_duplicate_number_count", len(dup_nums))
    print("template_exact_mismatch_count", len(exact_mismatch))
    print("project_mode_sample_ch5_hits", len(project_mode_sample_ch5_hits))
    print("project_domain_leak_hits", len(project_domain_leak_hits))
    print("missing_figure_captions", len(missing_fig_caps))
    print("missing_table_captions", len(missing_tab_caps))
    print("fig_caption", fig_caption)
    print("figures_inserted", figures_inserted)
    print("figures_pack_exported", figures_pack_exported)
    print("diagram_generation_gate_fail", diagram_generation_gate_fail)
    print("ordered_pack_src_exists", int(src.exists()))
    print("ordered_pack_dst_exists", int(dst.exists()))
    print("ordered_pack_src_bmp_count", len(src_bmp))
    print("ordered_pack_dst_bmp_count", len(dst_bmp))
    print("ordered_pack_src_map_exists", int(src_map))
    print("ordered_pack_dst_map_exists", int(dst_map))
    print("ordered_pack_seq_ok", int(seq_ok))
    print("ordered_pack_gate_fail", ordered_pack_gate_fail)
    print("chapter5_fake_inline_rows", len(chapter5_fake_inline_rows))
    print("three_line_table_violations", len(three_line_table_violations))
    print("tables_with_tblStyle", tables_with_tblStyle)
    print("cell_vertical_override_violations", len(cell_vertical_override))
    print("cell_three_line_shape_violations", len(cell_three_line_shape))

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
