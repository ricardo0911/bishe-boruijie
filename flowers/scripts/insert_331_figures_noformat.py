from __future__ import annotations

import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Cm


DOCX_NAME = "花店智能管理系统-毕业论文-初稿.docx"
EDGE_EXE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")

FIGS = [
    (
        "图3-10 订单明细实体属性描述",
        "图3-10_订单明细实体属性描述.png",
        "order_item",
        ["id", "order_id", "product_id", "product_title", "unit_price", "quantity", "line_amount", "created_at", "updated_at"],
    ),
    (
        "图3-11 库存锁流水实体属性描述",
        "图3-11_库存锁流水实体属性描述.png",
        "stock_lock",
        ["id", "order_id", "order_no", "flower_id", "batch_id", "lock_qty", "status", "expires_at", "created_at", "updated_at"],
    ),
    (
        "图3-12 支付流水实体属性描述",
        "图3-12_支付流水实体属性描述.png",
        "payment_log",
        ["id", "order_id", "order_no", "transaction_id", "payment_channel", "pay_amount", "result_code", "notify_time", "created_at"],
    ),
    (
        "图3-13 推荐结果实体属性描述",
        "图3-13_推荐结果实体属性描述.png",
        "recommendation_result",
        ["id", "user_id", "product_id", "score", "reason", "generated_at"],
    ),
]


def build_single_entity_html(entity: str, attrs: list[str], caption: str) -> str:
    attrs_js = ", ".join([f"'{x}'" for x in attrs])
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{caption}</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
  font-family: "SimSun", "宋体", serif;
  background: #fff;
  width: 760px;
  padding: 40px 40px 20px;
}}
.wrap {{ display:flex; flex-direction:column; align-items:center; gap:14px; }}
.entity-container {{ position: relative; width: 500px; height: 460px; }}
.entity-container svg {{
  position: absolute; top: 0; left: 0; width: 500px; height: 460px; z-index: 1;
}}
.entity-center {{
  position: absolute; width: 132px; height: 62px; background: #2f5d87; color:#fff;
  display:flex; align-items:center; justify-content:center; font-size:16px; font-weight:bold;
  z-index:3; transform: translate(-50%, -50%);
}}
.attr-oval {{
  position:absolute; padding:6px 14px; background:#2f5d87; color:#fff; border-radius:50%;
  font-size:13px; white-space:nowrap; z-index:3; transform: translate(-50%, -50%); text-align:center;
}}
.caption {{ text-align:center; font-size:14px; color:#222; letter-spacing:2px; }}
</style>
</head>
<body>
<div class="wrap">
  <div id="container" class="entity-container"></div>
  <div class="caption">{caption}</div>
</div>
<script>
(() => {{
  const CX=250, CY=230, R=168;
  const attrs=[{attrs_js}];
  const n=attrs.length;
  const container=document.getElementById('container');
  const svg=document.createElementNS('http://www.w3.org/2000/svg','svg');
  svg.setAttribute('viewBox','0 0 500 460');
  container.appendChild(svg);

  const center=document.createElement('div');
  center.className='entity-center';
  center.textContent='{entity}';
  center.style.left=CX+'px';
  center.style.top=CY+'px';
  container.appendChild(center);

  for(let i=0;i<n;i++) {{
    const angle=(2*Math.PI*i/n)-Math.PI/2;
    const x=CX+R*Math.cos(angle), y=CY+R*Math.sin(angle);
    const line=document.createElementNS('http://www.w3.org/2000/svg','line');
    line.setAttribute('x1',CX); line.setAttribute('y1',CY);
    line.setAttribute('x2',x); line.setAttribute('y2',y);
    line.setAttribute('stroke','#2f5d87'); line.setAttribute('stroke-width','1');
    svg.appendChild(line);
    const oval=document.createElement('div');
    oval.className='attr-oval';
    oval.textContent=attrs[i];
    oval.style.left=x+'px'; oval.style.top=y+'px';
    container.appendChild(oval);
  }}
}})();
</script>
</body>
</html>"""


def has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def ensure_331_order(doc: Document) -> None:
    from patch_thesis_331_expand import ADDED_LINES  # local script in same folder

    remove_prefixes = (
        "（6）订单明细实体",
        "图3-10 订单明细实体属性描述",
        "（7）库存锁流水实体",
        "图3-11 库存锁流水实体属性描述",
        "（8）支付流水实体",
        "图3-12 支付流水实体属性描述",
        "（9）推荐结果实体",
        "图3-13 推荐结果实体属性描述",
    )
    for p in list(doc.paragraphs):
        t = (p.text or "").strip()
        if any(t.startswith(x) for x in remove_prefixes):
            p._element.getparent().remove(p._element)

    anchor = None
    for p in doc.paragraphs:
        if (p.text or "").strip().startswith("3.3.2 逻辑结构设计"):
            anchor = p
            break
    if anchor is None:
        raise RuntimeError("未找到 3.3.2 逻辑结构设计 段落")

    for line in ADDED_LINES:
        anchor.insert_paragraph_before(line)


def render_figures(workdir: Path) -> dict[str, Path]:
    if not EDGE_EXE.exists():
        raise FileNotFoundError(f"未找到 Edge: {EDGE_EXE}")

    out_dir = workdir / "图包_补充实体图"
    out_dir.mkdir(exist_ok=True)

    rendered: dict[str, Path] = {}
    with tempfile.TemporaryDirectory(prefix="entity331_") as td:
        tmp = Path(td)
        for caption, png_name, entity, attrs in FIGS:
            html = build_single_entity_html(entity, attrs, caption)
            html_path = tmp / (png_name.replace(".png", ".html"))
            html_path.write_text(html, encoding="utf-8")

            png_tmp = tmp / png_name
            cmd = [
                str(EDGE_EXE),
                "--headless",
                "--disable-gpu",
                "--hide-scrollbars",
                "--window-size=900,700",
                f"--screenshot={png_tmp}",
                html_path.as_uri(),
            ]
            subprocess.run(cmd, check=True)

            png_out = out_dir / png_name
            png_out.write_bytes(png_tmp.read_bytes())
            rendered[caption] = png_out
    return rendered


def insert_images(doc: Document, mapping: dict[str, Path]) -> tuple[int, int]:
    inserted = 0
    skipped = 0
    paragraphs = doc.paragraphs

    for i, p in enumerate(paragraphs):
        t = (p.text or "").strip()
        if t not in mapping:
            continue
        if i > 0 and has_drawing(paragraphs[i - 1]):
            skipped += 1
            continue
        img_p = p.insert_paragraph_before("")
        run = img_p.add_run()
        run.add_picture(str(mapping[t]), width=Cm(13))
        img_p.alignment = 1
        inserted += 1

    return inserted, skipped


def main() -> int:
    root = Path.cwd()
    docx_path = root / DOCX_NAME
    if not docx_path.exists():
        raise FileNotFoundError(f"未找到文档: {docx_path}")

    doc = Document(docx_path)

    # Ensure section 3.3.1 new entities are in correct order.
    ensure_331_order(doc)

    # Render four new images.
    mapping = render_figures(root)

    # Insert images in place.
    inserted, skipped = insert_images(doc, mapping)
    doc.save(docx_path)
    print(f"inserted={inserted} skipped={skipped}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
