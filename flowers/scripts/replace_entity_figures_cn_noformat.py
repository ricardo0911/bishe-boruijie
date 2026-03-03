from __future__ import annotations

import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Cm


DOCX_NAME = "花店智能管理系统-毕业论文-初稿.docx"
EDGE_EXE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")

FIG_SPECS = [
    (
        "图3-8 库存批次实体属性描述",
        "图3-8_库存批次实体属性描述.png",
        "库存批次",
        ["批次编号", "花材编号", "供应商", "到货时间", "枯萎时间", "原始库存", "当前库存", "锁定库存", "单位成本", "质量等级", "创建时间"],
    ),
    (
        "图3-9 售后记录实体属性描述",
        "图3-9_售后记录实体属性描述.png",
        "售后记录",
        ["售后编号", "退款单号", "订单编号", "退款金额", "退款原因", "申请说明", "处理状态", "申请时间", "审核时间", "退款时间", "交易流水"],
    ),
    (
        "图3-10 订单明细实体属性描述",
        "图3-10_订单明细实体属性描述.png",
        "订单明细",
        ["明细编号", "订单编号", "商品编号", "商品名称", "单价", "数量", "小计金额", "创建时间", "更新时间"],
    ),
    (
        "图3-11 库存锁流水实体属性描述",
        "图3-11_库存锁流水实体属性描述.png",
        "库存锁流水",
        ["流水编号", "订单编号", "订单号", "花材编号", "批次编号", "锁定数量", "锁定状态", "过期时间", "创建时间", "更新时间"],
    ),
    (
        "图3-12 支付流水实体属性描述",
        "图3-12_支付流水实体属性描述.png",
        "支付流水",
        ["流水编号", "订单编号", "订单号", "交易号", "支付渠道", "支付金额", "结果编码", "回调时间", "创建时间"],
    ),
    (
        "图3-13 推荐结果实体属性描述",
        "图3-13_推荐结果实体属性描述.png",
        "推荐结果",
        ["结果编号", "用户编号", "商品编号", "推荐分值", "推荐理由", "生成时间"],
    ),
]


def build_entity_html(center_label: str, attrs: list[str], caption: str) -> str:
    attrs_js = ", ".join([f"'{x}'" for x in attrs])
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{caption}</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
  font-family: "SimSun", "宋体", serif;
  background: #ffffff;
  width: 760px;
  padding: 34px 34px 16px;
}}
.wrap {{ display: flex; flex-direction: column; align-items: center; gap: 10px; }}
.entity-container {{ position: relative; width: 620px; height: 520px; }}
.entity-container svg {{
  position: absolute; top: 0; left: 0; width: 620px; height: 520px; z-index: 1;
}}
.entity-center {{
  position: absolute; width: 176px; height: 86px;
  background: #184f86; color: #fff;
  display: flex; align-items: center; justify-content: center;
  font-size: 36px; font-weight: bold; letter-spacing: 2px;
  z-index: 3; transform: translate(-50%, -50%);
}}
.attr-oval {{
  position: absolute;
  min-width: 86px;
  padding: 8px 18px;
  background: #184f86; color: #fff;
  border-radius: 999px;
  font-size: 16px;
  white-space: nowrap;
  text-align: center;
  z-index: 3;
  transform: translate(-50%, -50%);
}}
.caption {{
  text-align: center;
  font-size: 15px;
  color: #222;
  letter-spacing: 1px;
}}
</style>
</head>
<body>
<div class="wrap">
  <div id="container" class="entity-container"></div>
  <div class="caption">{caption}</div>
</div>
<script>
(() => {{
  const CX = 310, CY = 250, R = 192;
  const attrs = [{attrs_js}];
  const n = attrs.length;
  const container = document.getElementById('container');
  const svg = document.createElementNS('http://www.w3.org/2000/svg', 'svg');
  svg.setAttribute('viewBox', '0 0 620 520');
  container.appendChild(svg);

  const center = document.createElement('div');
  center.className = 'entity-center';
  center.textContent = '{center_label}';
  center.style.left = CX + 'px';
  center.style.top = CY + 'px';
  container.appendChild(center);

  for (let i = 0; i < n; i++) {{
    const angle = (2 * Math.PI * i / n) - Math.PI / 2;
    const x = CX + R * Math.cos(angle);
    const y = CY + R * Math.sin(angle);

    const line = document.createElementNS('http://www.w3.org/2000/svg', 'line');
    line.setAttribute('x1', CX);
    line.setAttribute('y1', CY);
    line.setAttribute('x2', x);
    line.setAttribute('y2', y);
    line.setAttribute('stroke', '#7ea2c8');
    line.setAttribute('stroke-width', '2');
    svg.appendChild(line);

    const oval = document.createElement('div');
    oval.className = 'attr-oval';
    oval.textContent = attrs[i];
    oval.style.left = x + 'px';
    oval.style.top = y + 'px';
    container.appendChild(oval);
  }}
}})();
</script>
</body>
</html>"""


def has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def render_pngs(workdir: Path) -> dict[str, Path]:
    if not EDGE_EXE.exists():
        raise FileNotFoundError(f"未找到 Edge 可执行文件：{EDGE_EXE}")

    out_dir = workdir / "图包_补充实体图"
    out_dir.mkdir(exist_ok=True)
    mapping: dict[str, Path] = {}

    with tempfile.TemporaryDirectory(prefix="cn_entity_figs_") as td:
        tmp = Path(td)
        for caption, filename, center, attrs in FIG_SPECS:
            html = build_entity_html(center, attrs, caption)
            html_file = tmp / (filename.replace(".png", ".html"))
            html_file.write_text(html, encoding="utf-8")
            png_file = tmp / filename

            cmd = [
                str(EDGE_EXE),
                "--headless",
                "--disable-gpu",
                "--hide-scrollbars",
                "--window-size=980,760",
                f"--screenshot={png_file}",
                html_file.as_uri(),
            ]
            subprocess.run(cmd, check=True)

            target = out_dir / filename
            target.write_bytes(png_file.read_bytes())
            mapping[caption] = target
    return mapping


def replace_in_doc(doc_path: Path, image_map: dict[str, Path]) -> tuple[int, int]:
    doc = Document(doc_path)
    inserted = 0
    removed = 0

    i = 0
    while i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        text = (p.text or "").strip()
        if text not in image_map:
            i += 1
            continue

        # Remove all immediate drawing paragraphs above caption (usually one).
        while i > 0 and has_drawing(doc.paragraphs[i - 1]):
            prev = doc.paragraphs[i - 1]
            prev._element.getparent().remove(prev._element)
            removed += 1
            i -= 1

        img_p = doc.paragraphs[i].insert_paragraph_before("")
        run = img_p.add_run()
        run.add_picture(str(image_map[text]), width=Cm(13))
        img_p.alignment = 1
        inserted += 1
        i += 1

    doc.save(doc_path)
    return inserted, removed


def main() -> int:
    root = Path.cwd()
    doc_path = root / DOCX_NAME
    if not doc_path.exists():
        raise FileNotFoundError(f"未找到文档：{doc_path}")

    mapping = render_pngs(root)
    inserted, removed = replace_in_doc(doc_path, mapping)
    print(f"inserted={inserted} removed_old={removed}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
