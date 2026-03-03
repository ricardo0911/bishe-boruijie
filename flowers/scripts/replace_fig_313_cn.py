from __future__ import annotations

import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Cm


DOCX = Path("花店智能管理系统-毕业论文-初稿.docx")
EDGE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")
CAPTION = "图3-13 推荐结果实体属性描述"


def build_html() -> str:
    attrs = ["结果编号", "用户编号", "商品编号", "推荐分值", "推荐理由", "生成时间"]
    attrs_js = ", ".join([f"'{x}'" for x in attrs])
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
  font-family: "SimSun", "宋体", serif;
  background: #fff;
  width: 760px;
  padding: 30px;
}}
.stage {{ position: relative; width: 620px; height: 500px; margin: 0 auto; }}
svg {{ position: absolute; left: 0; top: 0; width: 620px; height: 500px; }}
.center {{
  position: absolute; left: 310px; top: 240px; transform: translate(-50%, -50%);
  width: 180px; height: 86px; background: #184f86; color: #fff;
  display: flex; align-items: center; justify-content: center;
  font-size: 38px; font-weight: bold; letter-spacing: 2px;
}}
.oval {{
  position: absolute; transform: translate(-50%, -50%);
  background: #184f86; color: #fff; border-radius: 999px;
  font-size: 16px; padding: 8px 16px; white-space: nowrap;
}}
</style>
</head>
<body>
<div class="stage" id="stage"></div>
<script>
(() => {{
  const CX=310, CY=240, R=185;
  const attrs=[{attrs_js}], n=attrs.length;
  const stage=document.getElementById('stage');
  const svg=document.createElementNS('http://www.w3.org/2000/svg','svg');
  svg.setAttribute('viewBox','0 0 620 500');
  stage.appendChild(svg);
  const c=document.createElement('div');
  c.className='center';
  c.textContent='推荐结果';
  stage.appendChild(c);
  for(let i=0;i<n;i++) {{
    const a=(2*Math.PI*i/n)-Math.PI/2;
    const x=CX+R*Math.cos(a), y=CY+R*Math.sin(a);
    const line=document.createElementNS('http://www.w3.org/2000/svg','line');
    line.setAttribute('x1',CX); line.setAttribute('y1',CY);
    line.setAttribute('x2',x); line.setAttribute('y2',y);
    line.setAttribute('stroke','#7ea2c8'); line.setAttribute('stroke-width','2');
    svg.appendChild(line);
    const o=document.createElement('div');
    o.className='oval'; o.textContent=attrs[i];
    o.style.left=x+'px'; o.style.top=y+'px';
    stage.appendChild(o);
  }}
}})();
</script>
</body>
</html>"""


def is_drawing(p) -> bool:
    return bool(p._p.xpath(".//w:drawing"))


def render_png(out_png: Path) -> None:
    if not EDGE.exists():
        raise FileNotFoundError(EDGE)
    with tempfile.TemporaryDirectory(prefix="fig313cn_") as td:
        td_path = Path(td)
        html_file = td_path / "fig313.html"
        html_file.write_text(build_html(), encoding="utf-8")
        cmd = [
            str(EDGE),
            "--headless",
            "--disable-gpu",
            "--hide-scrollbars",
            "--window-size=980,760",
            f"--screenshot={out_png}",
            html_file.as_uri(),
        ]
        subprocess.run(cmd, check=True)


def replace_figure(doc_path: Path, png: Path) -> None:
    d = Document(doc_path)
    for i, p in enumerate(d.paragraphs):
        t = (p.text or "").strip()
        if t != CAPTION:
            continue
        # Remove contiguous drawing paragraphs immediately above caption.
        j = i - 1
        while j >= 0 and is_drawing(d.paragraphs[j]):
            target = d.paragraphs[j]
            target._element.getparent().remove(target._element)
            j -= 1
            i -= 1
        # Insert new image above caption.
        img_p = d.paragraphs[i].insert_paragraph_before("")
        img_p.alignment = 1
        img_p.add_run().add_picture(str(png), width=Cm(13))
        break
    d.save(doc_path)


def main() -> int:
    root = Path.cwd()
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    out_dir = root / "图包_补充实体图"
    out_dir.mkdir(exist_ok=True)
    png = out_dir / "图3-13_推荐结果实体属性描述_中文版.png"
    render_png(png)
    replace_figure(DOCX, png)
    print(f"replaced {CAPTION} with {png.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
