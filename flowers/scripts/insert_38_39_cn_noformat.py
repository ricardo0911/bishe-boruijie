from __future__ import annotations

import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Cm


DOCX_NAME = "花店智能管理系统-毕业论文-初稿.docx"
EDGE_EXE = Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe")

SPECS = [
    (
        "图3-8 库存批次实体属性描述",
        "图3-8_库存批次实体属性描述.png",
        "库存批次",
        ["批次编号", "花材编号", "供应商", "到货时间", "枯萎时间", "原始库存", "当前库存", "锁定库存", "单位成本", "质量等级", "创建时间"],
    ),
    (
        "图3-9 售后记录实体属性描述",
        "图3-9_售后记录实体属性描述.png",
        "售后记录",
        ["售后编号", "退款单号", "订单编号", "退款金额", "退款原因", "申请说明", "处理状态", "申请时间", "审核时间", "退款时间", "交易流水"],
    ),
]


def html(center: str, attrs: list[str], caption: str) -> str:
    items = ", ".join([f"'{x}'" for x in attrs])
    return f"""<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0">
<style>
* {{ margin:0; padding:0; box-sizing:border-box; }}
body {{ font-family: "SimSun","宋体",serif; background:#fff; width:760px; padding:30px; }}
.c {{ position:relative; width:620px; height:520px; margin:0 auto; }}
svg {{ position:absolute; left:0; top:0; width:620px; height:520px; }}
.center {{ position:absolute; left:310px; top:250px; transform:translate(-50%,-50%); width:176px; height:86px; background:#184f86; color:#fff; font-size:36px; font-weight:bold; display:flex; align-items:center; justify-content:center; }}
.a {{ position:absolute; transform:translate(-50%,-50%); background:#184f86; color:#fff; border-radius:999px; font-size:16px; padding:8px 18px; white-space:nowrap; }}
</style></head><body>
<div class="c" id="c"></div>
<script>
(() => {{
 const CX=310,CY=250,R=192,attrs=[{items}],n=attrs.length;
 const c=document.getElementById('c');
 const s=document.createElementNS('http://www.w3.org/2000/svg','svg'); s.setAttribute('viewBox','0 0 620 520'); c.appendChild(s);
 const ce=document.createElement('div'); ce.className='center'; ce.textContent='{center}'; c.appendChild(ce);
 for(let i=0;i<n;i++) {{
  const ang=(2*Math.PI*i/n)-Math.PI/2;
  const x=CX+R*Math.cos(ang), y=CY+R*Math.sin(ang);
  const l=document.createElementNS('http://www.w3.org/2000/svg','line');
  l.setAttribute('x1',CX); l.setAttribute('y1',CY); l.setAttribute('x2',x); l.setAttribute('y2',y);
  l.setAttribute('stroke','#7ea2c8'); l.setAttribute('stroke-width','2'); s.appendChild(l);
  const a=document.createElement('div'); a.className='a'; a.textContent=attrs[i]; a.style.left=x+'px'; a.style.top=y+'px'; c.appendChild(a);
 }}
}})();
</script>
</body></html>"""


def has_drawing(p) -> bool:
    return bool(p._p.xpath(".//w:drawing"))


def main() -> int:
    root = Path.cwd()
    out = root / "图包_补充实体图"
    out.mkdir(exist_ok=True)

    if not EDGE_EXE.exists():
        raise FileNotFoundError(EDGE_EXE)

    with tempfile.TemporaryDirectory(prefix="fig389_") as td:
        td_path = Path(td)
        generated = {}
        for cap, name, center, attrs in SPECS:
            h = td_path / f"{name[:-4]}.html"
            p = td_path / name
            h.write_text(html(center, attrs, cap), encoding="utf-8")
            cmd = [
                str(EDGE_EXE),
                "--headless",
                "--disable-gpu",
                "--hide-scrollbars",
                "--window-size=980,760",
                f"--screenshot={p}",
                h.as_uri(),
            ]
            subprocess.run(cmd, check=True)
            target = out / name
            target.write_bytes(p.read_bytes())
            generated[cap] = target

    doc_path = root / DOCX_NAME
    d = Document(doc_path)
    inserted = 0
    for i, p in enumerate(d.paragraphs):
        t = (p.text or "").strip()
        if t not in generated:
            continue
        if i > 0 and has_drawing(d.paragraphs[i - 1]):
            continue
        img_p = p.insert_paragraph_before("")
        img_p.alignment = 1
        img_p.add_run().add_picture(str(generated[t]), width=Cm(13))
        inserted += 1

    d.save(doc_path)
    print(f"inserted={inserted}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
